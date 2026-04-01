"""
Report generator — produces a markdown summary and a formatted Word document.
"""
import json
import pathlib
from datetime import datetime


def save_report(report_data: dict, output_dir: str = ".") -> dict:
    """
    Save the shortlist report as both markdown and Word docx.
    report_data keys:
      - job_title: str
      - shortlist: list of candidate dicts
      - similar_profiles: list of profile dicts
      - summary: str (executive summary text)
    Returns: { markdown_path, docx_path }
    """
    output = pathlib.Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    job_title_slug = report_data.get("job_title", "role").replace(" ", "_")[:30]
    base_name = f"Shortlist_{job_title_slug}_{timestamp}"

    md_path = output / f"{base_name}.md"
    docx_path = output / f"{base_name}.docx"

    md_content = _build_markdown(report_data)
    md_path.write_text(md_content, encoding="utf-8")

    _build_docx(report_data, md_content, docx_path)

    return {"markdown_path": str(md_path), "docx_path": str(docx_path)}


def _build_markdown(data: dict) -> str:
    lines = []
    job_title = data.get("job_title", "Role")
    now = datetime.now().strftime("%B %d, %Y")

    lines += [
        f"# Candidate Shortlist — {job_title}",
        f"*Generated: {now}*",
        "",
        "---",
        "",
        "## Executive Summary",
        "",
        data.get("summary", ""),
        "",
        "---",
        "",
        "## Shortlisted Candidates",
        "",
    ]

    for rank, c in enumerate(data.get("shortlist", []), 1):
        lines += [
            f"### {rank}. {c.get('name', 'Unknown')} — Score: {c.get('score', 'N/A')}/100",
            "",
            f"**CV File:** {c.get('cv_file', '')}",
            f"**Current Role:** {c.get('current_role', '')}",
            f"**LinkedIn:** {c.get('linkedin_url', 'Not found')}",
            f"**LinkedIn Confidence:** {c.get('linkedin_confidence', 'N/A')}",
            "",
            "**Strengths:**",
        ]
        for s in c.get("strengths", []):
            lines.append(f"- {s}")
        lines += ["", "**Gaps / Watch Points:**"]
        for g in c.get("gaps", []):
            lines.append(f"- {g}")
        lines += [
            "",
            "**LinkedIn Validation:**",
            c.get("linkedin_validation", "Not validated"),
            "",
            "---",
            "",
        ]

    # Similar profiles
    similar = data.get("similar_profiles", [])
    if similar:
        lines += [
            "## Similar Profiles Found on LinkedIn",
            "*Candidates not in the CV pool who may be worth approaching.*",
            "",
        ]
        for p in similar:
            lines += [
                f"**{p.get('name', 'Unknown')}**",
                f"- Headline: {p.get('headline', '')}",
                f"- Location: {p.get('location', '')}",
                f"- LinkedIn: {p.get('url', '')}",
                f"- Source: {p.get('source', '')}",
                "",
            ]
        lines += ["---", ""]

    lines += [
        "## Recommended Next Steps",
        "",
        data.get("next_steps", ""),
        "",
    ]

    return "\n".join(lines)


def _build_docx(data: dict, md_content: str, output_path: pathlib.Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    NAVY  = (0x1A, 0x1A, 0x2E)
    SLATE = (0x2C, 0x3E, 0x50)
    RED   = (0xC0, 0x39, 0x2B)
    GRAY  = (0x7F, 0x8C, 0x8D)
    BLUE  = (0x1A, 0x52, 0x76)
    GREEN = (0x1E, 0x8B, 0x4C)

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2.5)

    def add_red_rule(para):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '1'); b.set(qn('w:color'), 'C0392B')
        pBdr.append(b); pPr.append(pBdr)

    def add_left_border(para, color='1A5276'):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        l = OxmlElement('w:left')
        l.set(qn('w:val'), 'single'); l.set(qn('w:sz'), '12')
        l.set(qn('w:space'), '4'); l.set(qn('w:color'), color)
        pBdr.append(l); pPr.append(pBdr)

    def set_shading(para, hex_color):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color); pPr.append(shd)

    def heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level > 1 else 0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text.upper() if level >= 2 else text)
        r.bold = True
        r.font.name = 'Calibri'
        if level == 1:
            r.font.size = Pt(20); r.font.color.rgb = RGBColor(*NAVY)
            add_red_rule(p)
        elif level == 2:
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(*RED)
        elif level == 3:
            r.font.size = Pt(12); r.font.color.rgb = RGBColor(*NAVY)
            r.text = text  # don't uppercase candidate names

    def body(text, color=SLATE, size=10.5, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)
        r.font.name = 'Calibri'; r.italic = italic
        return p

    def bullet(text, color=SLATE):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(text)
        r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(*color)
        r.font.name = 'Calibri'

    def label_value(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True; r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(*SLATE); r1.font.name = 'Calibri'
        r2 = p.add_run(str(value))
        r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(*SLATE)
        r2.font.name = 'Calibri'

    def score_badge(score):
        """Score as a highlighted inline label."""
        try:
            s = int(score)
            color = '1E8B4C' if s >= 75 else 'E67E22' if s >= 55 else 'C0392B'
        except Exception:
            color = '7F8C8D'
        return color

    # ── Document ──────────────────────────────────────────────────────────────
    job_title = data.get("job_title", "Role")
    heading(f"Candidate Shortlist — {job_title}", level=1)
    body(f"Generated: {datetime.now().strftime('%B %d, %Y')}  |  Confidential — HR Use Only",
         color=GRAY, size=9, italic=True)

    # Executive Summary
    heading("Executive Summary", level=2)
    p = body(data.get("summary", ""))
    set_shading(p, 'EBF5FB')
    add_left_border(p, '1A5276')

    # Shortlist
    heading("Shortlisted Candidates", level=2)

    for rank, c in enumerate(data.get("shortlist", []), 1):
        name = c.get("name", "Unknown")
        score = c.get("score", "N/A")
        badge_color = score_badge(score)

        # Candidate header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"#{rank}  {name}")
        r1.bold = True; r1.font.size = Pt(13)
        r1.font.color.rgb = RGBColor(*NAVY); r1.font.name = 'Calibri'
        r2 = p.add_run(f"   {score}/100")
        r2.bold = True; r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(*tuple(int(badge_color[i:i+2], 16) for i in (0,2,4)))
        r2.font.name = 'Calibri'

        label_value("CV File", c.get("cv_file", ""))
        label_value("Current Role", c.get("current_role", ""))
        label_value("LinkedIn", c.get("linkedin_url", "Not found"))
        label_value("LinkedIn Match", c.get("linkedin_confidence", "N/A"))

        # Strengths
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(6)
        r = p2.add_run("Strengths")
        r.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(*GREEN); r.font.name = 'Calibri'
        for s in c.get("strengths", []):
            bullet(s, color=SLATE)

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(4)
        r = p3.add_run("Gaps / Watch Points")
        r.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(*RED); r.font.name = 'Calibri'
        for g in c.get("gaps", []):
            bullet(g, color=SLATE)

        # LinkedIn validation
        validation = c.get("linkedin_validation", "")
        if validation:
            p4 = body(f"LinkedIn Validation: {validation}", color=(0x5D, 0x6D, 0x7E), italic=True)
            add_left_border(p4, '1A5276')

        doc.add_paragraph()  # spacer

    # Similar profiles
    similar = data.get("similar_profiles", [])
    if similar:
        heading("Similar Profiles on LinkedIn", level=2)
        body("The following profiles were not in the CV pool but match the role criteria and may be worth approaching.",
             color=GRAY, size=10, italic=True)
        for p_data in similar:
            p = doc.add_paragraph()
            r = p.add_run(p_data.get("name", "Unknown"))
            r.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(*NAVY); r.font.name = 'Calibri'
            label_value("Headline", p_data.get("headline", ""))
            label_value("Location", p_data.get("location", ""))
            label_value("LinkedIn", p_data.get("url", ""))
            doc.add_paragraph()

    # Next steps
    heading("Recommended Next Steps", level=2)
    next_steps = data.get("next_steps", "")
    if next_steps:
        for line in next_steps.split("\n"):
            line = line.strip().lstrip("-•").strip()
            if line:
                bullet(line)

    doc.save(str(output_path))
