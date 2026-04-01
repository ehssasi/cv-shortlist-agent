"""
CV Parser — extracts text from PDF and DOCX files.
"""
import os
import pathlib


def parse_cv(file_path: str) -> dict:
    """
    Parse a CV file (PDF or DOCX) and return extracted text + metadata.
    Returns: { filename, text, error }
    """
    path = pathlib.Path(file_path)
    result = {"filename": path.name, "file_path": str(path), "text": "", "error": None}

    if not path.exists():
        result["error"] = f"File not found: {file_path}"
        return result

    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            result["text"] = _parse_pdf(path)
        elif suffix in (".docx", ".doc"):
            result["text"] = _parse_docx(path)
        elif suffix == ".txt":
            result["text"] = path.read_text(encoding="utf-8", errors="ignore")
        else:
            result["error"] = f"Unsupported file type: {suffix}"
    except Exception as e:
        result["error"] = str(e)

    return result


def _parse_pdf(path: pathlib.Path) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _parse_docx(path: pathlib.Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def list_cvs(folder_path: str) -> list[dict]:
    """
    List all CV files in a folder. Returns list of { filename, file_path }.
    Supported: .pdf, .docx, .doc, .txt
    """
    folder = pathlib.Path(folder_path)
    if not folder.exists():
        return []

    supported = {".pdf", ".docx", ".doc", ".txt"}
    files = []
    for f in sorted(folder.iterdir()):
        if f.is_file() and f.suffix.lower() in supported:
            files.append({"filename": f.name, "file_path": str(f)})
    return files
